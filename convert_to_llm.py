import argparse
import json
import os
import subprocess
import sys

def ensure_dependencies():
    try:
        import docx
        import bs4
    except ImportError:
        print("Installing required dependencies (python-docx, beautifulsoup4)...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "beautifulsoup4"])

ensure_dependencies()

from docx import Document
from bs4 import BeautifulSoup

def process_telegram(json_path, out_path):
    doc = Document()
    doc.add_heading('Telegram Chat', level=1)
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    messages = data.get('messages', [])
    for msg in messages:
        if msg.get('type') != 'message':
            continue
            
        date = msg.get('date', '')
        sender = msg.get('from', 'Unknown')
        
        text_entities = msg.get('text_entities', [])
        text = ""
        for entity in text_entities:
            # standard text or custom emoji
            text += entity.get('text', '')
            
        if not text:
            text = ""
            
        attachments = []
        if 'photo' in msg:
            attachments.append('[Photo]')
        if 'media_type' in msg:
            mt = msg['media_type']
            if mt == 'voice_message':
                attachments.append('[Voice Message]')
            elif mt == 'sticker':
                attachments.append('[Sticker]')
            elif mt == 'video_file':
                attachments.append('[Video]')
            elif mt == 'video_message':
                attachments.append('[Video Message]')
            else:
                attachments.append(f"[{mt}]")
        elif 'file' in msg:
            attachments.append('[File]')
            
        content = text
        if attachments:
            content += " " + " ".join(attachments)
            
        content = content.strip()
        if content:
            p = doc.add_paragraph()
            p.add_run(f"[{date}] {sender}: ").bold = True
            p.add_run(content)
            
    doc.save(out_path)

def process_discord(html_path, out_path):
    doc = Document()
    doc.add_heading('Discord Chat', level=1)
    
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
        
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Discord messages are list items
    messages = soup.find_all('li', class_=lambda x: x and 'messageListItem_' in x)
    current_sender = "Unknown"
    current_date = ""
    
    for msg in messages:
        header = msg.find('h3', class_=lambda x: x and 'header_' in x)
        if header:
            user_span = header.find('span', class_=lambda x: x and 'username_' in x)
            if user_span:
                current_sender = user_span.get('data-text', user_span.text.strip())
            time_tag = header.find('time')
            if time_tag:
                current_date = time_tag.get('datetime', time_tag.text.strip())
        else:
            time_tag = msg.find('time')
            if time_tag:
                current_date = time_tag.get('datetime', current_date)
             
        content_div = msg.find('div', class_=lambda x: x and 'messageContent_' in x)
        text = ""
        if content_div:
            # Handle emojis
            for img in content_div.find_all('img'):
                alt = img.get('alt')
                if alt:
                    img.replace_with(alt)
                else:
                    img.replace_with('[Emoji]')
            text = content_div.get_text(separator='', strip=True)
            
        accessories = msg.find('div', id=lambda x: x and x.startswith('message-accessories-'))
        attachments = []
        if accessories:
            imgs = accessories.find_all('img')
            for img in imgs:
                alt = img.get('alt', '')
                if 'Стикер' in alt or 'Sticker' in alt or 'sticker' in img.get('src', '').lower():
                    attachments.append('[Sticker]')
                elif 'avatar' not in img.get('class', []) and 'emoji' not in img.get('class', []):
                    attachments.append('[Image]')
                    
            if accessories.find('video'):
                attachments.append('[Video]')
            
            # audio
            if accessories.find('audio'):
                attachments.append('[Voice Message]')
        
        content = text
        if attachments:
            # Remove duplicated [Sticker]
            attachments = list(dict.fromkeys(attachments))
            content += " " + " ".join(attachments)
             
        content = content.strip()
        if content:
            p = doc.add_paragraph()
            # Clean up discord datetimes directly if possible, else leave as is
            clean_date = current_date.replace('T', ' ')[:19]
            p.add_run(f"[{clean_date}] {current_sender}: ").bold = True
            p.add_run(content)
        
    doc.save(out_path)

def main():
    parser = argparse.ArgumentParser(
        description="Convert Telegram (JSON) and Discord (HTML) chat exports to .docx for LLM tools."
    )
    parser.add_argument(
        "export_dir",
        nargs="?",
        help="Directory containing result.json and/or discord-chat.html",
    )
    parser.add_argument("--telegram", metavar="PATH", help="Path to Telegram result.json")
    parser.add_argument("--telegram-out", metavar="PATH", help="Output path for Telegram .docx")
    parser.add_argument("--discord", metavar="PATH", help="Path to Discord discord-chat.html")
    parser.add_argument("--discord-out", metavar="PATH", help="Output path for Discord .docx")
    args = parser.parse_args()

    if args.export_dir:
        base_dir = os.path.abspath(args.export_dir)
        tg_in = args.telegram or os.path.join(base_dir, "result.json")
        tg_out = args.telegram_out or os.path.join(base_dir, "telegram_chat.docx")
        discord_in = args.discord or os.path.join(base_dir, "discord-chat.html")
        discord_out = args.discord_out or os.path.join(base_dir, "discord_chat.docx")
    elif args.telegram or args.discord:
        tg_in = args.telegram
        tg_out = args.telegram_out or (
            os.path.splitext(tg_in)[0] + ".docx" if tg_in else None
        )
        discord_in = args.discord
        discord_out = args.discord_out or (
            os.path.splitext(discord_in)[0] + ".docx" if discord_in else None
        )
    else:
        parser.error("Provide export_dir or --telegram / --discord")

    processed = False

    if tg_in:
        print("Processing Telegram...")
        if os.path.exists(tg_in):
            process_telegram(tg_in, tg_out)
            print("Telegram saved to", tg_out)
            processed = True
        else:
            print("Telegram JSON not found:", tg_in)

    if discord_in:
        print("Processing Discord...")
        if os.path.exists(discord_in):
            process_discord(discord_in, discord_out)
            print("Discord saved to", discord_out)
            processed = True
        else:
            print("Discord HTML not found:", discord_in)

    if not processed:
        sys.exit(1)


if __name__ == "__main__":
    main()
